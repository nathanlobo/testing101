"""
McLaren CreditMemo Agent v1.0
A Streamlit application for automated credit memo generation from financial PDFs
Hackathon: Orix McLaren 2026
"""

import streamlit as st
import fitz  # PyMuPDF
import os
from pathlib import Path
from groq import Groq
import plotly.graph_objects as go
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Local persistence paths (keeps uploads and memos between sessions)
DATA_DIR = Path("data")
PDF_DIR = DATA_DIR / "pdf_uploads"
MEMO_DIR = DATA_DIR / "memos"

# Page config
st.set_page_config(
    page_title="McLaren CreditMemo Agent v1.0",
    page_icon="🏎️",
    layout="wide"
)

# System prompt for the AI
SYSTEM_PROMPT = """You are the McLaren CreditMemo Agent. Your job is to convert raw PDF text into a professional Credit Memo.

Rules:
- EXTRACT facts accurately. Cite page numbers for every claim (e.g., [Page 4]).
- STRUCTURE the output with Markdown headers: # Executive Summary, # Key Metrics, # Risk Analysis.
- ASSIGN Confidence Tags: ✅ (Direct Match), ⚠️ (Inferred), ❌ (Missing).
- TONE: Senior Financial Analyst. Concise.
- If data is missing for a specific metric (e.g., Churn), explicitly state 'Data Gap'.

Output Format:
# Executive Summary
- [Bullet 1 with citation]
- [Bullet 2 with citation]
- [Bullet 3 with citation]
- [Bullet 4 with citation]
- [Bullet 5 with citation]

# Key Metrics
| Metric | Current Year | Prior Year | Change | Confidence |
|--------|-------------|-----------|--------|-----------|
| Revenue | $XX.XM | $XX.XM | +X% | ✅/⚠️/❌ |
| EBITDA | $XX.XM | $XX.XM | +X% | ✅/⚠️/❌ |
| Debt Ratio | X.XX | X.XX | +X% | ✅/⚠️/❌ |

# Risk Analysis
## Risk 1: [Risk Name] [Page X]
**Description:** [Brief description]
**Mitigation:** [Suggestion]

## Risk 2: [Risk Name] [Page X]
**Description:** [Brief description]
**Mitigation:** [Suggestion]

## Risk 3: [Risk Name] [Page X]
**Description:** [Brief description]
**Mitigation:** [Suggestion]
"""

# Initialize session state
if 'memo_content' not in st.session_state:
    st.session_state.memo_content = None
if 'pdf_text' not in st.session_state:
    st.session_state.pdf_text = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'pdf_doc' not in st.session_state:
    st.session_state.pdf_doc = None
if 'pdf_bytes' not in st.session_state:
    st.session_state.pdf_bytes = None
if 'last_saved' not in st.session_state:
    st.session_state.last_saved = None


def extract_pdf_text_with_pages(pdf_file):
    """Extract text from PDF with page number tracking"""
    pdf_bytes = pdf_file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    full_text = ""
    page_texts = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        page_texts.append((page_num + 1, text))
        full_text += f"\n\n--- PAGE {page_num + 1} ---\n{text}"
    
    # Store PDF for viewer
    st.session_state.pdf_doc = pdf_bytes
    st.session_state.pdf_bytes = pdf_bytes
    
    return full_text, page_texts, doc


def extract_pdf_text_from_bytes(pdf_bytes: bytes):
    """Extract text with page markers from raw PDF bytes"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    full_text = ""
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        full_text += f"\n\n--- PAGE {page_num + 1} ---\n{text}"
    return full_text, doc


def generate_credit_memo(pdf_text, groq_api_key):
    """Generate credit memo using Groq API"""
    try:
        client = Groq(api_key=groq_api_key)
        
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": f"Analyze this financial document and generate a Credit Memo:\n\n{pdf_text[:15000]}"  # Limit text to avoid token limits
                }
            ],
            model="llama-3.1-8b-instant",
            temperature=0.3,
            max_tokens=2000
        )
        
        return chat_completion.choices[0].message.content
    except Exception as e:
        return f"Error generating memo: {str(e)}"


def refine_memo(current_memo, pdf_text, refinement_request, groq_api_key):
    """Refine existing memo based on user request"""
    try:
        client = Groq(api_key=groq_api_key)
        
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": SYSTEM_PROMPT
                },
                {
                    "role": "user",
                    "content": f"Original PDF text (reference):\n{pdf_text[:8000]}\n\nCurrent Memo:\n{current_memo}\n\nUser Request: {refinement_request}\n\nPlease update the memo accordingly while maintaining the same structure and citation format."
                }
            ],
            model="llama-3.1-8b-instant",
            temperature=0.3,
            max_tokens=2000
        )
        
        return chat_completion.choices[0].message.content
    except Exception as e:
        return f"Error refining memo: {str(e)}"


def create_trend_chart():
    """Create a sample trend chart for metrics visualization"""
    fig = go.Figure()
    
    # Sample data
    years = ['2022', '2023', '2024']
    revenue = [150.5, 175.2, 198.7]
    ebitda = [35.2, 42.8, 51.3]
    
    fig.add_trace(go.Scatter(
        x=years, y=revenue,
        mode='lines+markers',
        name='Revenue ($M)',
        line=dict(color='#00D9FF', width=3),
        marker=dict(size=10)
    ))
    
    fig.add_trace(go.Scatter(
        x=years, y=ebitda,
        mode='lines+markers',
        name='EBITDA ($M)',
        line=dict(color='#FF6B00', width=3),
        marker=dict(size=10)
    ))
    
    fig.update_layout(
        title='Financial Trend Analysis',
        xaxis_title='Year',
        yaxis_title='Amount ($M)',
        template='plotly_dark',
        height=300,
        margin=dict(l=20, r=20, t=40, b=20)
    )
    
    return fig


def export_to_docx(memo_content):
    """Export memo to Word document"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Credit Memo', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Add content
    for line in memo_content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), 1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), 2)
        elif line.strip():
            doc.add_paragraph(line)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def ensure_storage_dirs():
    """Create local storage folders if missing"""
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    MEMO_DIR.mkdir(parents=True, exist_ok=True)


def persist_run(uploaded_file, memo_content):
    """Save uploaded PDF and generated memo to disk for persistence"""
    ensure_storage_dirs()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = Path(uploaded_file.name).name.replace(" ", "_")

    pdf_path = PDF_DIR / f"{timestamp}_{safe_name}"
    memo_path = MEMO_DIR / f"{timestamp}_{Path(safe_name).stem}.md"

    if st.session_state.pdf_bytes:
        pdf_path.write_bytes(st.session_state.pdf_bytes)
    memo_path.write_text(memo_content or "", encoding="utf-8")

    st.session_state.last_saved = {
        "pdf": str(pdf_path),
        "memo": str(memo_path),
        "timestamp": timestamp,
    }
    return pdf_path, memo_path


def list_saved_runs():
    """Return list of saved runs paired by timestamp"""
    ensure_storage_dirs()
    def ts_key(path: Path):
        parts = path.name.split("_")
        if len(parts) >= 2:
            return f"{parts[0]}_{parts[1]}"
        return parts[0]

    pdfs = {ts_key(p): p for p in PDF_DIR.glob("*.pdf")}
    memos = {ts_key(m): m for m in MEMO_DIR.glob("*.md")}
    runs = []
    for ts, memo_path in memos.items():
        pdf_path = pdfs.get(ts)
        if pdf_path:
            runs.append({
                "timestamp": ts,
                "memo": memo_path,
                "pdf": pdf_path
            })
    runs.sort(key=lambda r: r["timestamp"], reverse=True)
    return runs


def load_saved_run(run):
    """Load memo and PDF into session state"""
    pdf_bytes = run["pdf"].read_bytes()
    pdf_text, doc = extract_pdf_text_from_bytes(pdf_bytes)
    memo_text = run["memo"].read_text(encoding="utf-8")

    st.session_state.memo_content = memo_text
    st.session_state.pdf_text = pdf_text
    st.session_state.pdf_doc = pdf_bytes
    st.session_state.pdf_bytes = pdf_bytes
    st.session_state.chat_history = []
    st.session_state.last_saved = {
        "pdf": str(run["pdf"]),
        "memo": str(run["memo"]),
        "timestamp": run["timestamp"],
    }


def format_run_label(run):
    """Return a human-friendly label for a saved run"""
    ts_raw = run.get("timestamp", "")
    ts_human = ts_raw
    try:
        ts_human = datetime.strptime(ts_raw, "%Y%m%d_%H%M%S").strftime("%Y-%m-%d %H:%M")
    except Exception:
        pass
    pdf_name = run.get("pdf").name if run.get("pdf") else "uploaded.pdf"
    parts = pdf_name.split("_", 2)
    clean_name = pdf_name
    if len(parts) >= 3:
        clean_name = parts[2]
    elif len(parts) == 2:
        clean_name = parts[1]
    return f"{ts_human} | {clean_name}"


# Main UI
st.title("🏎️ McLaren CreditMemo Agent v1.0")
st.caption("AI-Powered Credit Memo Generation | Powered by Groq Llama-3")

# Sidebar for API key and upload
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # API Key input
    groq_api_key = st.text_input(
        "Groq API Key",
        type="password",
        value=st.secrets.get("GROQ_API_KEY", "") if "GROQ_API_KEY" in st.secrets else "",
        help="Enter your Groq API key or set it in .streamlit/secrets.toml"
    )
    
    st.divider()

    # Load previous runs
    st.header("📂 Load Saved Run")
    saved_runs = list_saved_runs()
    if saved_runs:
        options = {format_run_label(r): r for r in saved_runs}
        selected_label = st.selectbox(
            "Select a saved memo",
            options=list(options.keys())
        )
        if st.button("🔄 Load Selected", use_container_width=True):
            load_saved_run(options[selected_label])
            st.success("Loaded saved memo and PDF")
            st.rerun()
    else:
        st.caption("No saved runs yet. Generate a memo to create one.")
    
    # File uploader
    st.header("📄 Upload Document")
    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type=['pdf'],
        help="Upload a Balance Sheet, Annual Report, or 10-K filing"
    )
    
    if uploaded_file and groq_api_key:
        if st.button("🚀 Generate Credit Memo", type="primary", use_container_width=True):
            with st.spinner("Extracting PDF content..."):
                pdf_text, page_texts, doc = extract_pdf_text_with_pages(uploaded_file)
                st.session_state.pdf_text = pdf_text
            
            with st.spinner("🏎️ Generating memo at McLaren speed..."):
                memo = generate_credit_memo(pdf_text, groq_api_key)
                st.session_state.memo_content = memo
                st.session_state.chat_history = []
                # Persist PDF + memo locally for durability
                pdf_path, memo_path = persist_run(uploaded_file, memo)
            
            st.success("✅ Memo generated and saved locally!")
            st.caption(f"Saved: PDF → {pdf_path.name} | Memo → {memo_path.name} (data/)")
            st.rerun()
    
    st.divider()
    
    # Stats
    st.header("📊 Stats")
    if st.session_state.pdf_text:
        st.metric("PDF Pages", st.session_state.pdf_text.count("--- PAGE"))
        st.metric("Total Characters", len(st.session_state.pdf_text))
    if st.session_state.last_saved:
        st.text(f"Last saved memo: {Path(st.session_state.last_saved['memo']).name}")
        st.text(f"Last saved PDF: {Path(st.session_state.last_saved['pdf']).name}")
    
    st.divider()
    
    # Info
    st.info("💡 **Tip:** After generating the memo, use the chat to refine sections like 'Simplify Risk Analysis' or 'Add more detail to Executive Summary'")

# Main content area - Split view
if st.session_state.memo_content:
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📄 Source Document")
        
        if st.session_state.pdf_doc:
            # Display PDF (simplified viewer)
            st.info("PDF loaded successfully. Navigate through pages to verify citations.")
            
            # Page selector
            pdf_bytes = st.session_state.pdf_doc
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            page_num = st.number_input(
                "Page", 
                min_value=1, 
                max_value=len(doc), 
                value=1,
                step=1
            )
            
            # Render selected page as image
            page = doc[page_num - 1]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            st.image(img_data, caption=f"Page {page_num} of {len(doc)}", use_container_width=True)
        
        # Trend chart
        st.subheader("📈 Financial Trends")
        st.plotly_chart(create_trend_chart(), use_container_width=True)
    
    with col2:
        st.header("📝 Generated Credit Memo")
        
        # Display memo
        st.markdown(st.session_state.memo_content)
        
        st.divider()
        
        # Export buttons
        col_export1, col_export2 = st.columns(2)
        
        with col_export1:
            # Markdown export
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button(
                label="📥 Download as .md",
                data=st.session_state.memo_content,
                file_name=f"credit_memo_{timestamp}.md",
                mime="text/markdown",
                use_container_width=True
            )
        
        with col_export2:
            # Word export
            docx_buffer = export_to_docx(st.session_state.memo_content)
            st.download_button(
                label="📥 Download as .docx",
                data=docx_buffer,
                file_name=f"credit_memo_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        st.divider()
        
        # Chat refinement interface
        st.subheader("💬 Refine Memo")
        
        # Display chat history
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.write(message["content"])
        
        # Chat input
        if refinement_request := st.chat_input("Ask to refine the memo (e.g., 'Simplify the risk section')"):
            # Add user message to history
            st.session_state.chat_history.append({
                "role": "user",
                "content": refinement_request
            })
            
            with st.spinner("Refining memo..."):
                refined_memo = refine_memo(
                    st.session_state.memo_content,
                    st.session_state.pdf_text,
                    refinement_request,
                    groq_api_key
                )
                
                st.session_state.memo_content = refined_memo
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": "✅ Memo updated!"
                })
            
            st.rerun()

else:
    # Welcome screen
    st.markdown("""
    ## Welcome to McLaren CreditMemo Agent 🏎️
    
    ### What this does:
    Transform dense financial PDFs into structured, audit-ready Credit Memos in **seconds**.
    
    ### Features:
    - **📄 Split-Screen View:** PDF on left, generated memo on right
    - **🎯 Smart Extraction:** Auto-generates Executive Summary, Key Metrics, and Risk Analysis
    - **🔍 Source Traceability:** Every claim cites [Page X] for audit compliance
    - **💬 Interactive Refinement:** Chat to adjust tone, add detail, or simplify sections
    - **📊 Visual Analytics:** Trend charts for key financial metrics
    - **📥 Export Ready:** Download as Markdown or Word document
    
    ### How to use:
    1. Enter your Groq API key in the sidebar (or set in secrets)
    2. Upload a financial PDF (Balance Sheet, 10-K, Annual Report)
    3. Click "Generate Credit Memo"
    4. Review, refine via chat, and export!
    
    ---
    
    **Built for:** Orix McLaren Hackathon 2026  
    **Stack:** Streamlit + PyMuPDF + Groq (Llama-3) + Plotly  
    **Speed:** Sub-5-second generation ⚡
    """)
    
    # Sample use cases
    st.subheader("💡 Sample Refinement Commands:")
    st.code("""
    - "Simplify the risk analysis section"
    - "Add more detail to the executive summary"
    - "Find the source for the Debt Ratio metric"
    - "Make the tone more formal"
    - "Highlight liquidity concerns"
    """)

# Footer
st.divider()
st.caption("McLaren CreditMemo Agent v1.0 | Hackathon Edition | Powered by Groq Llama-3 🚀")
